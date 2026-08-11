import asyncio
import io
from uuid import UUID, uuid4

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from openpyxl import Workbook
from pypdf import PdfWriter
from xlwt import Workbook as XlsWorkbook

from app.api.dependencies import get_document_ingestion_service
from app.core.config import Settings
from app.core.errors import ApiError
from app.features.documents import parser
from app.features.documents.parser import limit_extracted_text, parse_document
from app.features.documents.schemas import DocumentKind
from app.features.documents.service import DocumentIngestionService, UploadInput
from app.main import app


class FakeDocumentStore:
    def __init__(self, workspace_id: UUID, fail_on_task: bool = False) -> None:
        self.workspace_id = workspace_id
        self.fail_on_task = fail_on_task
        self.uploaded_paths: list[str] = []
        self.uploaded_mime_types: list[str] = []
        self.documents: list[dict[str, object]] = []
        self.tasks: list[dict[str, object]] = []

    def workspace_exists(self, workspace_id: UUID) -> bool:
        return workspace_id == self.workspace_id

    def list_workspace_documents(self, workspace_id: UUID) -> list[dict[str, object]]:
        if workspace_id != self.workspace_id:
            return []
        return self.documents

    def upload_file(self, path: str, _content: bytes, mime_type: str) -> None:
        self.uploaded_paths.append(path)
        self.uploaded_mime_types.append(mime_type)

    def create_source_document(self, data: dict[str, object]) -> dict[str, object]:
        self.documents.append(data)
        return data

    def create_parse_task(self, data: dict[str, object]) -> dict[str, object]:
        if self.fail_on_task:
            raise RuntimeError("database unavailable")
        self.tasks.append(data)
        return data

    def delete_source_document(self, document_id: UUID) -> None:
        self.documents = [item for item in self.documents if item["id"] != str(document_id)]

    def delete_file(self, path: str) -> None:
        self.uploaded_paths.remove(path)


class FakeTaskQueue:
    def __init__(self) -> None:
        self.task_ids: list[UUID] = []

    async def enqueue_parse_document(self, task_id: UUID) -> None:
        self.task_ids.append(task_id)


class FailingUploadStore(FakeDocumentStore):
    def upload_file(self, _path: str, _content: bytes, _mime_type: str) -> None:
        raise ApiError(
            502,
            "DOCUMENT_STORAGE_UPLOAD_FAILED",
            "文件存储服务拒绝了上传请求，请检查文件格式后重试。",
        )


class FailingTaskQueue:
    async def enqueue_parse_document(self, task_id: UUID) -> None:
        raise RuntimeError(f"redis unavailable: {task_id}")


def test_parse_csv_normalizes_rows() -> None:
    result = parse_document("商品,价格\n保温杯,99".encode(), "text/csv", "products.csv")

    assert result.text == "商品 | 价格\n保温杯 | 99"
    assert result.metadata["rowCount"] == 2
    assert result.needs_ocr is False


def test_parse_markdown_is_treated_as_text() -> None:
    result = parse_document("# 标题\n\n正文".encode(), "text/markdown", "说明.md")

    assert result.text == "# 标题\n\n正文"
    assert result.metadata["parser"] == "text"


def test_parse_docx_extracts_paragraphs_and_tables() -> None:
    document = DocxDocument()
    document.add_paragraph("会议结论：继续推进。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "负责人"
    table.cell(0, 1).text = "张三"
    buffer = io.BytesIO()
    document.save(buffer)

    result = parse_document(
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "会议纪要.docx",
    )

    assert "会议结论：继续推进。" in result.text
    assert "负责人 | 张三" in result.text
    assert result.metadata["parser"] == "docx"


def test_limit_extracted_text_marks_truncated_content() -> None:
    parsed = parse_document(b"abcdefgh", "text/plain", "notes.txt")
    limited = limit_extracted_text(parsed, 4)

    assert limited.text == "abcd"
    assert limited.metadata["truncated"] is True
    assert limited.metadata["originalCharacterCount"] == 8


def test_parse_image_returns_metadata_and_ocr_marker() -> None:
    # A minimal 1x1 transparent PNG.
    image = bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
        "0000000d49444154789c6360000002000154a24f5d0000000049454e44ae426082"
    )

    result = parse_document(image, "image/png", "product.png")

    assert result.metadata["width"] == 1
    assert result.needs_ocr is True


def test_parse_xlsx_extracts_worksheet_rows() -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "竞品"
    worksheet.append(["商品", "价格"])
    worksheet.append(["竞品A", 129])
    buffer = io.BytesIO()
    workbook.save(buffer)

    result = parse_document(
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "competitors.xlsx",
    )

    assert result.text == "[竞品] 商品 | 价格\n[竞品] 竞品A | 129"
    assert result.metadata["worksheetNames"] == ["竞品"]


def test_parse_xls_extracts_worksheet_rows() -> None:
    workbook = XlsWorkbook()
    worksheet = workbook.add_sheet("竞品")
    worksheet.write(0, 0, "商品")
    worksheet.write(0, 1, "价格")
    worksheet.write(1, 0, "竞品A")
    worksheet.write(1, 1, 129)
    buffer = io.BytesIO()
    workbook.save(buffer)

    result = parse_document(buffer.getvalue(), "application/vnd.ms-excel", "competitors.xls")

    assert result.text == "[竞品] 商品 | 价格\n[竞品] 竞品A | 129.0"
    assert result.metadata["parser"] == "xls"
    assert result.metadata["worksheetNames"] == ["竞品"]


def test_parse_doc_extracts_legacy_word_text(monkeypatch: pytest.MonkeyPatch) -> None:
    class LegacyDocResult:
        text = "旧版 Word 内容"

    monkeypatch.setattr(parser, "extract_legacy_doc_text", lambda _content: LegacyDocResult())

    result = parse_document(b"legacy-doc", "application/msword", "meeting.doc")

    assert result.text == "旧版 Word 内容"
    assert result.metadata["parser"] == "doc"


def test_parse_empty_pdf_marks_document_for_ocr() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buffer = io.BytesIO()
    writer.write(buffer)

    result = parse_document(buffer.getvalue(), "application/pdf", "scan.pdf")

    assert result.text == ""
    assert result.metadata["pageCount"] == 1
    assert result.needs_ocr is True


def test_upload_creates_pending_document_and_parse_task() -> None:
    workspace_id = uuid4()
    fake_store = FakeDocumentStore(workspace_id)
    fake_queue = FakeTaskQueue()
    service = DocumentIngestionService(
        store=fake_store,
        queue=fake_queue,
        settings=Settings(document_max_upload_size_bytes=1_000_000),
    )
    app.dependency_overrides[get_document_ingestion_service] = lambda: service
    client = TestClient(app)

    try:
        response = client.post(
            f"/api/v1/workspaces/{workspace_id}/documents",
            data={"kind": "DATASET"},
            files={"file": ("competitor.csv", "商品,价格\n竞品A,129".encode(), "text/csv")},
        )
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 201
    payload = response.json()
    assert payload["workspace_id"] == str(workspace_id)
    assert payload["status"] == "PENDING"
    assert payload["task_status"] == "QUEUED"
    assert fake_store.uploaded_paths[0].startswith(f"{workspace_id}/")
    assert fake_store.documents[0]["storage_path"] == fake_store.uploaded_paths[0]
    assert fake_store.tasks[0]["task_type"] == "PARSE_DOCUMENT"
    assert fake_queue.task_ids == [UUID(payload["task_id"])]


def test_upload_rejects_an_unsupported_file_type() -> None:
    workspace_id = uuid4()
    fake_store = FakeDocumentStore(workspace_id)
    fake_queue = FakeTaskQueue()
    service = DocumentIngestionService(
        store=fake_store,
        queue=fake_queue,
        settings=Settings(document_max_upload_size_bytes=1_000_000),
    )
    app.dependency_overrides[get_document_ingestion_service] = lambda: service
    client = TestClient(app)

    try:
        response = client.post(
            f"/api/v1/workspaces/{workspace_id}/documents",
            data={"kind": "GENERAL"},
            files={"file": ("payload.exe", b"unsafe", "application/octet-stream")},
        )
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 415
    assert response.json()["error"]["code"] == "DOCUMENT_TYPE_UNSUPPORTED"
    assert fake_store.uploaded_paths == []


def test_upload_returns_a_clear_error_when_storage_rejects_file() -> None:
    workspace_id = uuid4()
    service = DocumentIngestionService(
        store=FailingUploadStore(workspace_id),
        queue=FakeTaskQueue(),
        settings=Settings(document_max_upload_size_bytes=1_000_000),
    )

    with pytest.raises(ApiError, match="文件存储服务拒绝了上传请求") as error:
        asyncio.run(
            service.ingest(
                UploadInput(
                    workspace_id=workspace_id,
                    kind=DocumentKind.GENERAL,
                    file_name="legacy.xls",
                    mime_type="application/vnd.ms-excel",
                    content=b"legacy-office-content",
                )
            )
        )

    assert error.value.code == "DOCUMENT_STORAGE_UPLOAD_FAILED"


def test_upload_normalizes_legacy_xls_mime_type_before_storing() -> None:
    workspace_id = uuid4()
    fake_store = FakeDocumentStore(workspace_id)
    service = DocumentIngestionService(
        store=fake_store,
        queue=FakeTaskQueue(),
        settings=Settings(document_max_upload_size_bytes=1_000_000),
    )

    asyncio.run(
        service.ingest(
            UploadInput(
                workspace_id=workspace_id,
                kind=DocumentKind.GENERAL,
                file_name="legacy.xls",
                mime_type="application/octet-stream",
                content=b"legacy-office-content",
            )
        )
    )

    assert fake_store.uploaded_mime_types == ["application/vnd.ms-excel"]
    assert fake_store.documents[0]["mime_type"] == "application/vnd.ms-excel"


def test_upload_keeps_unicode_file_name_and_cleans_up_on_persistence_failure() -> None:
    workspace_id = uuid4()
    fake_store = FakeDocumentStore(workspace_id, fail_on_task=True)
    service = DocumentIngestionService(
        store=fake_store,
        queue=FakeTaskQueue(),
        settings=Settings(document_max_upload_size_bytes=1_000_000),
    )

    with pytest.raises(RuntimeError, match="database unavailable"):
        asyncio.run(
            service.ingest(
                upload=UploadInput(
                    workspace_id=workspace_id,
                    kind=DocumentKind.NOTE,
                    file_name="会议纪要.md",
                    mime_type="text/markdown",
                    content=b"# notes",
                )
            )
        )

    assert fake_store.uploaded_paths == []
    assert fake_store.documents == []


def test_upload_keeps_queued_task_when_queue_is_temporarily_unavailable() -> None:
    workspace_id = uuid4()
    fake_store = FakeDocumentStore(workspace_id)
    service = DocumentIngestionService(
        store=fake_store,
        queue=FailingTaskQueue(),
        settings=Settings(document_max_upload_size_bytes=1_000_000),
    )

    response = asyncio.run(
        service.ingest(
            upload=UploadInput(
                workspace_id=workspace_id,
                kind=DocumentKind.NOTE,
                file_name="会议纪要.md",
                mime_type="text/markdown",
                content=b"# notes",
            )
        )
    )

    assert response.file_name == "会议纪要.md"
    assert fake_store.tasks[0]["status"] == "QUEUED"


def test_list_documents_returns_workspace_materials() -> None:
    workspace_id = uuid4()
    fake_store = FakeDocumentStore(workspace_id)
    fake_store.documents.append(
        {
            "id": str(uuid4()),
            "workspace_id": str(workspace_id),
            "kind": "DATASET",
            "file_name": "reviews.csv",
            "mime_type": "text/csv",
            "size_bytes": 128,
            "status": "READY",
            "error_message": None,
            "metadata": {"parse": {"parser": "csv"}},
            "created_at": "2026-08-06T00:00:00Z",
            "updated_at": "2026-08-06T00:00:00Z",
        }
    )
    service = DocumentIngestionService(
        store=fake_store,
        queue=FakeTaskQueue(),
        settings=Settings(document_max_upload_size_bytes=1_000_000),
    )
    app.dependency_overrides[get_document_ingestion_service] = lambda: service
    client = TestClient(app)

    try:
        response = client.get(f"/api/v1/workspaces/{workspace_id}/documents")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    assert response.json()[0]["file_name"] == "reviews.csv"
