from __future__ import annotations

import csv
import io
from dataclasses import dataclass
from pathlib import PurePath

import xlrd
from docx import Document as DocxDocument
from legacy_doc import extract_text as extract_legacy_doc_text
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader

from app.core.errors import ApiError

TEXT_MIME_TYPES = {"text/plain", "text/markdown", "text/csv"}
SPREADSHEET_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME_TYPE = "application/msword"
PDF_MIME_TYPE = "application/pdf"
IMAGE_MIME_TYPES = {"image/jpeg", "image/png", "image/webp"}
SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".docx",
    ".doc",
    ".xlsx",
    ".xls",
    ".pdf",
    ".jpg",
    ".jpeg",
    ".png",
    ".webp",
}
MIME_TYPE_BY_EXTENSION = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".csv": "text/csv",
    ".xls": "application/vnd.ms-excel",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".pdf": "application/pdf",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".webp": "image/webp",
}


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    metadata: dict[str, object]
    needs_ocr: bool = False


def normalize_document_mime_type(mime_type: str, filename: str) -> str:
    """Prefer a known file extension over browsers' generic upload MIME types."""
    return MIME_TYPE_BY_EXTENSION.get(PurePath(filename).suffix.lower(), mime_type)


def limit_extracted_text(parsed: ParsedDocument, max_characters: int) -> ParsedDocument:
    """Cap extracted text before storing derived artifacts or sending it to an indexer."""
    if max_characters <= 0:
        raise ValueError("max_characters must be positive")
    if len(parsed.text) <= max_characters:
        return parsed
    return ParsedDocument(
        text=parsed.text[:max_characters],
        metadata={
            **parsed.metadata,
            "characterCount": max_characters,
            "originalCharacterCount": len(parsed.text),
            "truncated": True,
        },
        needs_ocr=parsed.needs_ocr,
    )


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue

    raise ApiError(
        422, "DOCUMENT_ENCODING_UNSUPPORTED", "文件编码无法识别，请使用 UTF-8 或 GB18030。"
    )


def _parse_text(content: bytes, mime_type: str) -> ParsedDocument:
    text = _decode_text(content)

    if mime_type != "text/csv":
        return ParsedDocument(text=text, metadata={"parser": "text", "characterCount": len(text)})

    rows = list(csv.reader(io.StringIO(text)))
    normalized_rows = [" | ".join(cell.strip() for cell in row if cell.strip()) for row in rows]
    extracted_text = "\n".join(row for row in normalized_rows if row).strip()
    return ParsedDocument(
        text=extracted_text,
        metadata={"parser": "csv", "rowCount": len(rows), "characterCount": len(extracted_text)},
    )


def _is_cjk_character(character: str) -> bool:
    return "\u3400" <= character <= "\u9fff" or "\uf900" <= character <= "\ufaff"


def _is_readable_spreadsheet_text(text: str) -> bool:
    """Recognize human-readable spreadsheet values after normalizing whitespace."""
    visible_characters = [character for character in text if not character.isspace()]
    if len(visible_characters) < 32:
        return True

    basic_character_count = sum(
        character.isascii() or _is_cjk_character(character) for character in visible_characters
    )
    # Chinese workbook prose is mostly CJK plus ASCII. A large proportion of
    # other glyphs is characteristic of raw device frames decoded as text.
    if basic_character_count / len(visible_characters) < 0.82:
        return False

    letters = [character.casefold() for character in visible_characters if character.isalpha()]
    highest_letter_ratio = (
        max(letters.count(character) for character in set(letters)) / len(letters)
        if letters
        else 0
    )
    return highest_letter_ratio <= 0.45


def _spreadsheet_cell_text(value: object) -> str | None:
    """Return readable spreadsheet text and reject embedded binary payloads.

    Some legacy XLS files place device frames or other raw bytes in a cell.
    xlrd can decode those bytes into Latin-1 and control characters, but that
    result is not meaningful RAG material and can poison an otherwise useful
    workbook's index.
    """
    readable_lines: list[str] = []
    for raw_line in str(value).replace("\r", "\n").split("\n"):
        control_character_count = sum(
            not character.isprintable() and character != "\t" for character in raw_line
        )
        # A long line with control bytes is a device frame, not formatted cell
        # text. Skip the line while retaining legitimate lines in the same cell.
        if len(raw_line) >= 32 and control_character_count >= 2:
            continue
        normalized = "".join(
            character if character.isprintable() or character == "\t" else " "
            for character in raw_line
        )
        text = " ".join(normalized.split())
        if text and _is_readable_spreadsheet_text(text):
            readable_lines.append(text)

    return " ".join(readable_lines) or None


def _parse_spreadsheet(content: bytes) -> ParsedDocument:
    try:
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as error:
        raise ApiError(
            422, "DOCUMENT_PARSE_FAILED", "无法读取 Excel 文件，请确认文件未损坏。"
        ) from error

    rows: list[str] = []
    worksheet_names: list[str] = []
    row_count = 0
    skipped_cell_count = 0

    for worksheet in workbook.worksheets:
        worksheet_names.append(worksheet.title)
        for row in worksheet.iter_rows(values_only=True):
            values: list[str] = []
            for value in row:
                if value is None or not str(value).strip():
                    continue
                extracted = _spreadsheet_cell_text(value)
                if extracted is None:
                    skipped_cell_count += 1
                    continue
                values.append(extracted)
            row_text = " | ".join(values)
            if row_text and _is_readable_spreadsheet_text(row_text):
                rows.append(f"[{worksheet.title}] {row_text}")
                row_count += 1
            elif row_text:
                skipped_cell_count += len(values)

    workbook.close()
    text = "\n".join(rows).strip()
    return ParsedDocument(
        text=text,
        metadata={
            "parser": "spreadsheet",
            "worksheetNames": worksheet_names,
            "rowCount": row_count,
            "skippedCellCount": skipped_cell_count,
            "characterCount": len(text),
        },
    )


def _parse_xls(content: bytes) -> ParsedDocument:
    try:
        workbook = xlrd.open_workbook(file_contents=content, on_demand=True)
    except Exception as error:
        raise ApiError(
            422, "DOCUMENT_PARSE_FAILED", "无法读取 XLS 文件，请确认文件未损坏。"
        ) from error

    rows: list[str] = []
    worksheet_names: list[str] = []
    row_count = 0
    skipped_cell_count = 0

    try:
        for worksheet in workbook.sheets():
            worksheet_names.append(worksheet.name)
            for row_index in range(worksheet.nrows):
                values: list[str] = []
                for column_index in range(worksheet.ncols):
                    value = worksheet.cell_value(row_index, column_index)
                    if not str(value).strip():
                        continue
                    extracted = _spreadsheet_cell_text(value)
                    if extracted is None:
                        skipped_cell_count += 1
                        continue
                    values.append(extracted)
                row_text = " | ".join(values)
                if row_text and _is_readable_spreadsheet_text(row_text):
                    rows.append(f"[{worksheet.name}] {row_text}")
                    row_count += 1
                elif row_text:
                    skipped_cell_count += len(values)
    finally:
        workbook.release_resources()

    text = "\n".join(rows).strip()
    return ParsedDocument(
        text=text,
        metadata={
            "parser": "xls",
            "worksheetNames": worksheet_names,
            "rowCount": row_count,
            "skippedCellCount": skipped_cell_count,
            "characterCount": len(text),
        },
    )


def _parse_docx(content: bytes) -> ParsedDocument:
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as error:
        raise ApiError(
            422, "DOCUMENT_PARSE_FAILED", "无法读取 DOCX 文件，请确认文件未损坏。"
        ) from error

    paragraphs = [
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()
            ]
            if cells:
                table_rows.append(" | ".join(cells))

    text = "\n\n".join(paragraphs + table_rows).strip()
    return ParsedDocument(
        text=text,
        metadata={
            "parser": "docx",
            "paragraphCount": len(paragraphs),
            "tableCount": len(document.tables),
            "tableRowCount": len(table_rows),
            "characterCount": len(text),
        },
    )


def _parse_doc(content: bytes) -> ParsedDocument:
    try:
        result = extract_legacy_doc_text(content)
        text = result.text.strip()
    except Exception as error:
        raise ApiError(
            422, "DOCUMENT_PARSE_FAILED", "无法读取 DOC 文件，请确认文件未损坏或未加密。"
        ) from error

    return ParsedDocument(
        text=text,
        metadata={
            "parser": "doc",
            "characterCount": len(text),
        },
    )


def _parse_pdf(content: bytes) -> ParsedDocument:
    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as error:
        raise ApiError(
            422, "DOCUMENT_PARSE_FAILED", "无法读取 PDF 文件，请确认文件未损坏或受密码保护。"
        ) from error

    text = "\n\n".join(page for page in pages if page).strip()
    return ParsedDocument(
        text=text,
        metadata={"parser": "pdf", "pageCount": len(pages), "characterCount": len(text)},
        needs_ocr=not bool(text),
    )


def _parse_image(content: bytes) -> ParsedDocument:
    try:
        image = Image.open(io.BytesIO(content))
        width, height = image.size
        image_format = image.format
    except Exception as error:
        raise ApiError(
            422, "DOCUMENT_PARSE_FAILED", "无法读取图片文件，请确认图片未损坏。"
        ) from error

    return ParsedDocument(
        text="",
        metadata={
            "parser": "image-metadata",
            "width": width,
            "height": height,
            "format": image_format,
        },
        needs_ocr=True,
    )


def validate_document_type(mime_type: str, filename: str) -> None:
    extension = PurePath(filename).suffix.lower()
    supported_mime_type = (
        mime_type in TEXT_MIME_TYPES
        or mime_type in SPREADSHEET_MIME_TYPES
        or mime_type == DOCX_MIME_TYPE
        or mime_type == DOC_MIME_TYPE
        or mime_type == PDF_MIME_TYPE
        or mime_type in IMAGE_MIME_TYPES
    )

    if extension and extension not in SUPPORTED_EXTENSIONS:
        raise ApiError(
            415,
            "DOCUMENT_TYPE_UNSUPPORTED",
            "仅支持 TXT、Markdown、CSV、XLS、XLSX、DOC、DOCX、PDF、JPG、PNG 和 WEBP 文件。",
        )
    if not supported_mime_type and extension not in SUPPORTED_EXTENSIONS:
        raise ApiError(
            415,
            "DOCUMENT_TYPE_UNSUPPORTED",
            "仅支持 TXT、Markdown、CSV、XLS、XLSX、DOC、DOCX、PDF、JPG、PNG 和 WEBP 文件。",
        )


def parse_document(content: bytes, mime_type: str, filename: str) -> ParsedDocument:
    """Extract deterministic text/metadata; OCR is deferred to the worker pipeline."""
    extension = PurePath(filename).suffix.lower()
    mime_type = normalize_document_mime_type(mime_type, filename)
    validate_document_type(mime_type, filename)

    if mime_type in TEXT_MIME_TYPES or extension in {".txt", ".md", ".markdown", ".csv"}:
        return _parse_text(content, "text/csv" if extension == ".csv" else mime_type)
    if extension == ".xls":
        return _parse_xls(content)
    if mime_type in SPREADSHEET_MIME_TYPES or extension == ".xlsx":
        return _parse_spreadsheet(content)
    if extension == ".doc":
        return _parse_doc(content)
    if mime_type == DOCX_MIME_TYPE or extension == ".docx":
        return _parse_docx(content)
    if mime_type == PDF_MIME_TYPE or extension == ".pdf":
        return _parse_pdf(content)
    if mime_type in IMAGE_MIME_TYPES or extension in {".jpg", ".jpeg", ".png", ".webp"}:
        return _parse_image(content)

    raise ApiError(
        415,
        "DOCUMENT_TYPE_UNSUPPORTED",
        "仅支持 TXT、Markdown、CSV、XLS、XLSX、DOC、DOCX、PDF、JPG、PNG 和 WEBP 文件。",
    )
